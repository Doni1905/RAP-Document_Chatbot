import os
from typing import List, Dict, Union
from pypdf import PdfReader
from docx import Document
import io

def load_documents(data_dir: str) -> List[Dict]:
    docs = []
    for fname in os.listdir(data_dir):
        fpath = os.path.join(data_dir, fname)
        if fname.lower().endswith(".pdf"):
            reader = PdfReader(fpath)
            total_pages = len(reader.pages)
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text:
                    docs.append({
                        "text": text,
                        "filename": fname,
                        "page": i + 1,
                        "total_pages": total_pages
                    })
        elif fname.lower().endswith(".docx"):
            doc = Document(fpath)
            paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
            docs.append({
                "text": "\n".join(paragraphs),
                "filename": fname,
                "paragraphs": paragraphs  # Keep for chunking
                # Do NOT set 'page' for DOCX
            })
    return docs


def load_uploaded_documents(uploaded_files: List) -> List[Dict]:
    """
    Load documents from Streamlit uploaded files
    """
    docs = []
    for uploaded_file in uploaded_files:
        fname = uploaded_file.name
        try:
            if fname.lower().endswith(".pdf"):
                # Read PDF from bytes
                pdf_bytes = uploaded_file.read()
                reader = PdfReader(io.BytesIO(pdf_bytes))
                total_pages = len(reader.pages)
                for i, page in enumerate(reader.pages):
                    text = page.extract_text()
                    if text:
                        docs.append({
                            "text": text,
                            "filename": fname,
                            "page": i + 1,
                            "total_pages": total_pages
                        })
            elif fname.lower().endswith(".docx"):
                # Read DOCX from bytes
                docx_bytes = uploaded_file.read()
                doc = Document(io.BytesIO(docx_bytes))
                paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
                docs.append({
                    "text": "\n".join(paragraphs),
                    "filename": fname,
                    "paragraphs": paragraphs
                })
            elif fname.lower().endswith(".txt"):
                # Read TXT from bytes
                txt_content = uploaded_file.read().decode('utf-8')
                docs.append({
                    "text": txt_content,
                    "filename": fname,
                    "paragraphs": txt_content.split('\n')
                })
        except Exception as e:
            print(f"Error processing {fname}: {str(e)}")
            continue
    return docs


def chunk_documents(docs: List[Dict], chunk_size: int = 500, overlap: int = 50) -> List[Dict]:
    chunks = []
    global_chunk_id = 0
    for doc in docs:
        if "paragraphs" in doc:  # DOCX
            words = []
            para_indices = []
            for i, para in enumerate(doc["paragraphs"]):
                para_words = para.split()
                words.extend(para_words)
                para_indices.extend([i+1] * len(para_words))  # 1-based paragraph numbers
            start = 0
            while start < len(words):
                end = min(start + chunk_size, len(words))
                chunk_words = words[start:end]
                chunk_text = " ".join(chunk_words)
                # Find paragraph range for this chunk
                para_range = set(para_indices[start:end])
                if len(para_range) == 1:
                    source_ref = f"Paragraph: {list(para_range)[0]}"
                else:
                    source_ref = f"Paragraphs: {min(para_range)}-{max(para_range)}"
                chunks.append({
                    "chunk_text": chunk_text,
                    "filename": doc["filename"],
                    "chunk_id": global_chunk_id,
                    "page": global_chunk_id + 1,  # Add this line for DOCX/TXT
                    "source_ref": source_ref
                    # Do NOT set 'total_pages' for DOCX
                })
                global_chunk_id += 1
                start += chunk_size - overlap
        else:
            # PDF logic
            text = doc["text"]
            words = text.split()
            start = 0
            while start < len(words):
                end = min(start + chunk_size, len(words))
                chunk_words = words[start:end]
                chunk_text = " ".join(chunk_words)
                chunks.append({
                    "chunk_text": chunk_text,
                    "filename": doc["filename"],
                    "page": doc["page"],
                    "total_pages": doc.get("total_pages"),
                    "chunk_id": global_chunk_id,
                    "source_ref": f"Page: {doc['page']}"
                })
                global_chunk_id += 1
                start += chunk_size - overlap
    return chunks